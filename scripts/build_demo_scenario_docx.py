from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/scenario_demo_assistant_copepodes.docx")

BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F7FA"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "687382"
DARK = "1F2933"
WHITE = "FFFFFF"
GOLD = "A66B00"
GREEN = "2F6B4F"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)


def set_font(run, size=11, bold=False, color=DARK, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text, bold=False, italic=False, color=DARK, size=11, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"{label} ")
    set_font(r, bold=True, color=BLUE)
    r = p.add_run(value)
    set_font(r)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"{label}\n")
    set_font(r, bold=True, color=accent)
    r = p.add_run(text)
    set_font(r, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_prompt(doc, text):
    add_callout(doc, "Prompt à envoyer à l’agent", text, fill="EEF5FB", accent=BLUE)


def add_step_header(doc, number, title, duration):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.page_break_before = True if number > 1 else False
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{number}. {title}")
    set_font(r, size=16, bold=True, color=BLUE)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    r = p2.add_run(f"Durée cible : {duration}")
    set_font(r, size=10, bold=True, color=MID_GRAY)


def add_demo_step(doc, number, title, duration, objective, say, prompt, expected, capability, caution, transition):
    add_step_header(doc, number, title, duration)
    add_label_value(doc, "Objectif.", objective)
    add_callout(doc, "Ce que tu dis", say, fill="F8F5EC", accent=GOLD)
    add_prompt(doc, prompt)
    add_label_value(doc, "Sortie attendue.", expected)
    add_label_value(doc, "Compétence mise en lumière.", capability)
    add_label_value(doc, "Vigilance scientifique.", caution)
    add_callout(doc, "Transition", transition, fill="F3F8F5", accent=GREEN)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for style_name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, "315D7A", 9, 4),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.text = "NeoLab | Démonstration de l’assistant scientifique copépodes"
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in hp.runs:
    set_font(run, size=9, color=MID_GRAY)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = fp.add_run("Scénario de démonstration")
set_font(run, size=9, color=MID_GRAY)

# Cover
add_text(doc, "GUIDE DE DÉMONSTRATION", bold=True, color=GOLD, size=11, after=12)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
r = p.add_run("Du fichier NeoLabs à une analyse zooplanctonique intégrée")
set_font(r, size=25, bold=True, color=BLUE)
add_text(
    doc,
    "Abondance, couverture d’échantillonnage, diversité, environnement CTD et ordination communauté-environnement",
    size=14,
    color="315D7A",
    after=18,
)

meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_widths(meta, [1.55, 4.95])
for row, (label, value) in enumerate([
    ("Public", "Professeur du laboratoire et professionnel de recherche"),
    ("Durée", "18 minutes, avec version courte de 10 minutes"),
    ("Fichier", "neolabs_taxonomy_abundance_amundsen_ctd.tsv"),
    ("Finalité", "Montrer une chaîne d’analyse scientifique crédible et utile à une demande de subvention"),
]):
    meta.cell(row, 0).text = label
    meta.cell(row, 1).text = value
    set_cell_shading(meta.cell(row, 0), LIGHT_BLUE)
    set_cell_shading(meta.cell(row, 1), WHITE)
    for cell in meta.rows[row].cells:
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for run in meta.cell(row, 0).paragraphs[0].runs:
        set_font(run, bold=True, color=BLUE)
    for run in meta.cell(row, 1).paragraphs[0].runs:
        set_font(run)

doc.add_paragraph()
add_callout(
    doc,
    "Message central",
    "L’assistant ne remplace pas l’interprétation scientifique. Il transforme un fichier réel en une analyse traçable : audit, contrôle qualité, diagnostic des lacunes, indicateurs biologiques, couplage environnemental et analyses multivariées exploratoires.",
    fill="EEF5FB",
)

doc.add_heading("Ce que la démonstration doit prouver", level=1)
for label, value in [
    ("Compréhension des données", "L’agent reconnaît les unités, le niveau taxonomique des lignes et les clés SAMPLE_ID + ANALYSIS_ID."),
    ("Rigueur", "Il sépare les données biologiques des variables CTD et filtre les jointures selon leur qualité."),
    ("Capacité analytique", "Il calcule abondance, saisonnalité, anomalies, diversité et composition communautaire."),
    ("Intégration", "Il relie les abondances aux variables température, salinité, oxygène, fluorescence et nitrate."),
    ("Valeur scientifique", "Il produit une ordination exploratoire et formule les lacunes qui justifient de nouvelles campagnes."),
]:
    add_label_value(doc, label + ".", value)

doc.add_heading("Repères issus du fichier réel", level=1)
rep = doc.add_table(rows=1, cols=3)
rep.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_widths(rep, [2.2, 1.3, 3.0])
for idx, text in enumerate(["Indicateur", "Valeur", "Lecture"]):
    rep.cell(0, idx).text = text
    set_cell_shading(rep.cell(0, idx), BLUE)
    for run in rep.cell(0, idx).paragraphs[0].runs:
        set_font(run, bold=True, color=WHITE)
for indicator, value, reading in [
    ("Lignes taxonomiques", "11 378", "Une ligne n’est pas un échantillon indépendant."),
    ("Sample-analysis", "421", "Niveau recommandé pour les analyses temporelles et CTD."),
    ("Stations", "242", "Couverture large mais effort très inégal."),
    ("Taxons", "275", "Assez riche pour diversité et ordination."),
    ("Période", "2010–2025", "Série longue mais discontinue."),
    ("CTD matched", "261 / 421", "Les analyses environnementales doivent filtrer les matches valides."),
]:
    cells = rep.add_row().cells
    for idx, text in enumerate([indicator, value, reading]):
        cells[idx].text = text
        set_cell_shading(cells[idx], WHITE if len(rep.rows) % 2 else LIGHT_GRAY)
        set_cell_margins(cells[idx])
        cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for run in cells[idx].paragraphs[0].runs:
            set_font(run, bold=(idx == 1), color=BLUE if idx == 1 else DARK)

doc.add_page_break()
doc.add_heading("Préparation avant la démonstration", level=1)
add_label_value(doc, "Fichier.", "Utiliser le TSV réel sans le modifier. Vérifier le chemin exact avant la séance.")
add_label_value(doc, "Session.", "Ouvrir une nouvelle conversation pour éviter toute table active héritée d’une autre démonstration.")
add_label_value(doc, "Réseau.", "Prévoir que Bio-ORACLE ou les services externes peuvent être plus lents ; le scénario principal repose d’abord sur le fichier local.")
add_label_value(doc, "Fenêtre.", "Afficher l’interface de l’agent et garder ce guide ouvert sur un deuxième écran.")
add_label_value(doc, "Règle.", "Ne pas annoncer un résultat avant qu’il soit calculé. Présenter les ordinations comme exploratoires.")

add_callout(
    doc,
    "Phrase d’ouverture",
    "Je vais partir d’un vrai fichier NeoLabs contenant des abondances taxonomiques déjà enrichies par des profils CTD Amundsen. L’objectif est de montrer comment l’assistant passe de données brutes à un diagnostic scientifique traçable, puis à des analyses écologiques classiques.",
    fill="F8F5EC",
    accent=GOLD,
)

add_demo_step(
    doc, 1, "Charger et comprendre le fichier", "1 min 30",
    "Montrer que l’agent inspecte la structure avant de calculer.",
    "Je commence par lui donner le fichier tel qu’il est produit. Je ne lui fournis ni dictionnaire manuel ni script préparé.",
    "Charge ce fichier et fais un audit scientifique initial. Identifie le nombre de lignes, le nombre de couples SAMPLE_ID + ANALYSIS_ID, les stations, les taxons, la période couverte, les unités d’abondance et les principales colonnes CTD. Ne fais encore aucune interprétation écologique.",
    "Résumé du fichier, reconnaissance de l’unité ind./m³, distinction entre lignes taxonomiques et échantillons, détection des colonnes CTD.",
    "Inspection automatique, compréhension des unités et identification du bon niveau d’analyse.",
    "Vérifier que l’agent ne traite pas les 11 378 lignes comme 11 378 prélèvements.",
    "Maintenant que le niveau des données est établi, je lui demande de vérifier si le fichier est suffisamment propre pour être interprété.",
)

add_demo_step(
    doc, 2, "Contrôle qualité des abondances et de la jointure CTD", "2 min",
    "Montrer que l’agent expose les limites au lieu de masquer les données imparfaites.",
    "Avant de chercher une tendance écologique, on vérifie les unités, les valeurs problématiques et la qualité du couplage avec Amundsen.",
    "Fais un contrôle qualité au niveau SAMPLE_ID + ANALYSIS_ID. Compare la disponibilité des abondances depth vol et flowmeter vol, détecte les valeurs négatives ou manquantes, résume ctd_match_status et décris les distributions de ctd_distance_km, ctd_time_delta_min et ctd_depth_coverage_m.",
    "Tableau QA/QC et graphique des statuts CTD : matched, hors plage, sans match, métadonnées manquantes.",
    "Traçabilité des calculs, gestion des unités et validation explicite des jointures multi-sources.",
    "Utiliser depth vol comme métrique principale ; flowmeter vol sert ici de contrôle. Ne pas mélanger les deux silencieusement.",
    "On sait maintenant quelles observations sont fiables. La prochaine question est : quand et où avons-nous réellement échantillonné ?",
)

add_demo_step(
    doc, 3, "Diagnostiquer la couverture temporelle", "2 min",
    "Quantifier les lacunes de suivi et la saisonnalité de l’effort d’échantillonnage.",
    "Une longue période ne signifie pas automatiquement une bonne série temporelle. Il faut regarder la régularité annuelle et mensuelle.",
    "Reconstruis sample_df avec une ligne par SAMPLE_ID + ANALYSIS_ID. Produis un histogramme du nombre d’échantillons par année et une heatmap année × mois. Signale les années et saisons absentes ou faiblement couvertes.",
    "Visualisation de 2010 à 2025, forte inégalité entre années, trou 2019–2022 et couverture limitée à juin–octobre.",
    "Reconstruction explicite de sample_df, analyse temporelle et diagnostic de lacunes.",
    "Ne pas interpréter une variation d’abondance annuelle sans tenir compte du nombre d’échantillons et des mois couverts.",
    "La série est longue, mais sa couverture est irrégulière. Je vais maintenant vérifier si cette irrégularité est aussi spatiale.",
)

add_demo_step(
    doc, 4, "Cartographier l’effort et les lacunes spatiales", "2 min",
    "Montrer la dimension géographique et distinguer présence de données et qualité CTD.",
    "Le modèle ne doit pas seulement dire combien de stations existent. Il doit montrer où elles se trouvent et où la couverture est faible.",
    "Crée une carte arctique des stations. La taille des points doit représenter le nombre de sample-analysis par station et la couleur le statut de couverture CTD. Ajoute une lecture rapide des zones ou stations peu répétées.",
    "Carte projetée avec côtes, stations très inégalement répétées et distinction matched/unmatched.",
    "Choix automatique de projection, agrégation spatiale et cartographie scientifique.",
    "Une station absente ne peut pas être inférée à partir du seul fichier ; parler de zones sous-échantillonnées parmi la couverture observée.",
    "Après avoir décrit l’effort, on peut enfin regarder le signal biologique en unités comparables.",
)

add_demo_step(
    doc, 5, "Décrire l’abondance et la composition", "2 min 30",
    "Produire les indicateurs biologiques les plus standards.",
    "Je demande d’abord les analyses les plus classiques : abondance normalisée, importance des copépodes et taxons dominants.",
    "À partir de Total abundance (ind./m3 depth vol), calcule par sample-analysis l’abondance totale zooplanctonique, l’abondance des copépodes et leur proportion. Montre ensuite les 15 taxons dominants en abondance cumulée et signale séparément les identifications génériques comme Animalia, Copepoda ou Calanus spp.",
    "Distribution très asymétrique, abondance médiane nettement inférieure aux maxima, taxons dominants tels que Oithona similis et Pseudocalanus spp., avec signalement des catégories génériques.",
    "Normalisation en ind./m³, filtrage taxonomique, agrégation et détection d’outliers.",
    "Utiliser une échelle logarithmique pour les distributions fortement asymétriques. Ne pas confondre dominance cumulée et fréquence d’occurrence.",
    "L’abondance montre combien d’organismes sont observés. La diversité montre comment cette abondance est répartie entre les taxons.",
)

add_demo_step(
    doc, 6, "Calculer richesse et diversité", "2 min",
    "Montrer une vraie analyse communautaire au niveau échantillon.",
    "On passe d’un simple classement de taxons à des indicateurs de structure communautaire.",
    "Construis une matrice sample-analysis × TAXON_ID avec les abondances en ind./m3. Calcule richesse taxonomique, Shannon, Simpson et Pielou. Compare leur distribution par année et identifie les échantillons très dominés ou particulièrement diversifiés.",
    "Tableau et graphiques des quatre indices, calculés sur la même matrice d’abondance.",
    "Pivot taxonomique, calculs écologiques standards et comparaison temporelle.",
    "Présenter les indices comme exploratoires : effort, résolution taxonomique et catégories génériques peuvent influencer la diversité.",
    "Nous avons maintenant la structure biologique. Je vais la replacer dans le contexte physico-chimique mesuré.",
)

add_demo_step(
    doc, 7, "Relier abondance et environnement CTD", "2 min 30",
    "Explorer les associations avec température, salinité, oxygène, fluorescence et nitrate.",
    "Cette étape montre que l’assistant ne joint pas aveuglément les sources : il limite l’analyse aux observations CTD valides.",
    "Garde uniquement ctd_match_status == matched. Trace log10(abondance totale + 1) en fonction de la température, de la salinité, de l’oxygène, de la fluorescence et du nitrate moyens sur l’intervalle de prélèvement. Colore les points par année et donne les corrélations exploratoires avec le nombre de paires utilisé.",
    "Cinq relations environnementales avec effectifs explicites ; associations globales probablement faibles ou hétérogènes.",
    "Filtrage QA, transformation logarithmique, analyses environnementales multi-variables et transparence sur les effectifs.",
    "Ne pas parler de causalité. Une corrélation faible peut refléter la diversité des régions, saisons, profondeurs ou protocoles.",
    "Les relations une variable à la fois restent limitées. Le moment fort consiste à examiner simultanément la communauté et l’environnement.",
)

add_demo_step(
    doc, 8, "Ordination communauté-environnement", "3 min",
    "Montrer la compétence analytique la plus avancée de la démonstration.",
    "Ici, on passe d’une série de graphes bivariés à une représentation multivariée de la structure communautaire.",
    "Réalise une analyse d’ordination exploratoire sur les échantillons CTD matched. Filtre les taxons présents dans moins de trois échantillons, applique une transformation Hellinger ou log1p, puis : 1) fais une PCA standardisée des variables CTD ; 2) fais une NMDS ou PCoA Bray-Curtis de la composition taxonomique ; 3) relie les axes aux variables température, salinité, oxygène, fluorescence et nitrate. Affiche le stress NMDS ou la variance expliquée et précise les limites.",
    "PCA environnementale, ordination Bray-Curtis colorée par température ou année, métrique de qualité et associations environnementales.",
    "Préparation de matrice communautaire, standardisation, distance Bray-Curtis, PCA/NMDS/PCoA et interprétation prudente.",
    "La RDA peut être proposée comme analyse exploratoire complémentaire. Ne pas annoncer une CCA ou une causalité sans méthode et tests formels validés.",
    "L’ordination ne donne pas une conclusion définitive ; elle révèle les gradients et groupes qui méritent une analyse confirmatoire et de nouveaux échantillonnages.",
)

add_demo_step(
    doc, 9, "Produire une synthèse orientée financement", "1 min 30",
    "Transformer les résultats en besoins de recherche concrets sans inventer de conclusion.",
    "Je termine en demandant au modèle de distinguer clairement ce que les données montrent, ce qu’elles suggèrent et ce qu’il faut financer.",
    "Rédige une synthèse d’une page structurée en trois parties : constats robustes, hypothèses exploratoires, priorités de collecte. Appuie chaque priorité sur une lacune mesurée : années ou saisons absentes, stations peu répétées, CTD non appariées, taxons ou gradients environnementaux à confirmer. Ne formule aucune causalité non démontrée.",
    "Argumentaire court reliant directement les lacunes observées à des campagnes, mesures ou analyses supplémentaires.",
    "Synthèse scientifique, traçabilité des preuves et préparation d’un livrable décisionnel.",
    "La justification financière doit découler des lacunes quantifiées, pas d’une promesse générale sur l’intelligence artificielle.",
    "La valeur du modèle est d’accélérer une chaîne complète et reproductible, tout en laissant au chercheur la validation et l’interprétation finale.",
)

doc.add_page_break()
doc.add_heading("Version courte — 10 minutes", level=1)
short = doc.add_table(rows=1, cols=3)
short.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_widths(short, [0.8, 2.0, 3.7])
for idx, text in enumerate(["Temps", "Étape", "Message"]):
    short.cell(0, idx).text = text
    set_cell_shading(short.cell(0, idx), BLUE)
    for run in short.cell(0, idx).paragraphs[0].runs:
        set_font(run, bold=True, color=WHITE)
for time, step, message in [
    ("1:00", "Audit", "Le modèle comprend unités, clés et niveau taxonomique."),
    ("1:30", "QA CTD", "Il expose les matches valides et les limites."),
    ("1:30", "Couverture", "Il quantifie les trous temporels et spatiaux."),
    ("2:00", "Abondance", "Il calcule ind./m³, copépodes et taxons dominants."),
    ("1:30", "Diversité", "Il reconstruit la matrice communauté."),
    ("2:00", "Ordination", "Il relie composition taxonomique et gradients CTD."),
    ("0:30", "Conclusion", "Il transforme les lacunes en priorités finançables."),
]:
    cells = short.add_row().cells
    for idx, text in enumerate([time, step, message]):
        cells[idx].text = text
        set_cell_shading(cells[idx], WHITE if len(short.rows) % 2 else LIGHT_GRAY)
        set_cell_margins(cells[idx])
        cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for run in cells[idx].paragraphs[0].runs:
            set_font(run, bold=(idx == 0), color=BLUE if idx == 0 else DARK)

doc.add_heading("Plan de repli si un calcul est lent", level=1)
add_label_value(doc, "Ordination.", "Afficher la PCA environnementale en priorité ; garder NMDS/PCoA comme deuxième sortie.")
add_label_value(doc, "Carte.", "Si Cartopy télécharge des fonds, passer à un tableau des stations les plus et moins répétées.")
add_label_value(doc, "Service externe.", "Ne pas dépendre de Bio-ORACLE pour le cœur de la démonstration. Le présenter comme extension future.")
add_label_value(doc, "Graphique.", "Si une sortie échoue, demander un tableau numérique équivalent plutôt que relancer plusieurs transformations à la fois.")

doc.add_page_break()
doc.add_heading("Phrases à éviter", level=1)
avoid = doc.add_table(rows=1, cols=2)
avoid.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_widths(avoid, [3.1, 3.4])
for idx, text in enumerate(["À éviter", "Formulation rigoureuse"]):
    avoid.cell(0, idx).text = text
    set_cell_shading(avoid.cell(0, idx), BLUE)
    for run in avoid.cell(0, idx).paragraphs[0].runs:
        set_font(run, bold=True, color=WHITE)
for bad, good in [
    ("La température contrôle l’abondance.", "L’association observée avec la température est exploratoire."),
    ("Nous avons 11 378 échantillons.", "Nous avons 11 378 lignes taxonomiques et 421 sample-analysis."),
    ("Toutes les données ont une CTD.", "261 sample-analysis ont un appariement CTD matched."),
    ("L’IA découvre les causes.", "L’assistant identifie des gradients et hypothèses à tester."),
    ("Il suffit de financer l’IA.", "Le financement vise les lacunes de collecte, validation et analyse révélées par les données."),
]:
    cells = avoid.add_row().cells
    for idx, text in enumerate([bad, good]):
        cells[idx].text = text
        set_cell_shading(cells[idx], WHITE if len(avoid.rows) % 2 else LIGHT_GRAY)
        set_cell_margins(cells[idx])
        for run in cells[idx].paragraphs[0].runs:
            set_font(run)

doc.add_heading("Conclusion orale recommandée", level=1)
add_callout(
    doc,
    "Conclusion",
    "Les analyses suivent les pratiques classiques de suivi zooplanctonique : abondance normalisée par volume filtré en ind./m³, abondance totale et copépodes, saisonnalité, diversité taxonomique, composition communautaire et mise en relation avec les variables CTD physico-chimiques. L’apport du modèle est d’accélérer cette chaîne tout en conservant les unités, les clés, la qualité des jointures et les limites. Il permet ainsi de transformer des lacunes mesurées en priorités scientifiques défendables pour une demande de subvention.",
    fill="EEF5FB",
)

doc.add_heading("Références méthodologiques utilisées pour cadrer la démo", level=1)
for source in [
    "OSPAR Quality Status Report 2023 — Changes in plankton biomass and abundance: https://oap.ospar.org/en/ospar-assessments/quality-status-reports/qsr-2023/indicator-assessments/changes-plankton-biomass-abundance/",
    "ICES Working Group on Zooplankton Ecology — synthèses saisonnières et séries de suivi: https://www.st.nmfs.noaa.gov/copepod/status-reports/crr307-wgze.pdf",
    "UNOLS — Collecting Zooplankton: https://www.unols.org/sites/default/files/Collecting_Zooplankton.pdf",
    "Journal of Plankton Research — analyses communauté zooplancton et variables environnementales: https://academic.oup.com/plankt/article-pdf/34/7/602/4279191/fbs029.pdf",
    "Diversity — exemple d’ordination RDA avec variables physico-chimiques: https://www.mdpi.com/1424-2818/11/11/203",
]:
    add_text(doc, source, size=9.5, color=MID_GRAY, after=5)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT.resolve())
